import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from io import BytesIO

# Estilos personalizados
st.markdown("""
    <style>
        /* Cambiar color de fondo y color de texto */
        .stApp {
            background-color: #000000;
            color: #FFFFFF;
        }
        
        /* Estilo para los encabezados */
        h1, h2, h3 {
            color: #4B8BBE;
            font-family: 'Times New Romans', sans-serif;
        }

        /* Estilo para el texto */
        p, li {
            font-family: 'Arial', sans-serif;
            color: #FFFFFF;
        }

        /* Estilo para el botón de descarga */
        .stDownloadButton {
            background-color: #FF0000;
            color: white;
            border-radius: 10px;
        }
        .stDownloadButton:hover {
            background-color: #FF6F61;
            opacity: 0.85;
        }

        /* Estilo para el botón de subir archivos */
        .stFileUploader {
            background-color: #FFD700;
            color: #333333;
            border-radius: 10px;
            border: 1px solid #FF6F61;
        }

        /* Estilo para el encabezado principal */
        .stTitle {
            color: #FF6F61;
            font-family: 'Georgia', serif;
        }
        
        /* Estilo para los subtítulos */
        .stHeader {
            color: #4B8BBE;
            font-family: 'Arial', sans-serif;
        }

    </style>
""", unsafe_allow_html=True)

# Función para generar el informe de Excel
def generar_informe_excel(pagos_vencidos_90_dias):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = f'CARTERA_DE_CLIENTES_VENCIDAS_{timestamp}.xlsx'
    
    # Crear un buffer en memoria
    buffer = BytesIO()
    
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        pagos_vencidos_90_dias.to_excel(writer, sheet_name='Vencidos_90_dias', index=False)
    
    st.success(f"Análisis completado. Resultados guardados en '{file_name}'.")
    
    # Colocar el buffer en la posición inicial
    buffer.seek(0)
    
    # Descargar el archivo de Excel
    st.download_button(label="Descargar archivo Excel", data=buffer, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# Función para extraer historial de clientes desde un archivo de Word
def extraer_historial_clientes(file):
    doc = Document(file)
    historial_clientes = [para.text for para in doc.paragraphs]
    return historial_clientes

# Función para generar el informe de Word
def generar_informe_word(pagos_vencidos_90_dias, historial_clientes, imagenes_papeles_paths, nombre_empresa, nombre_fraudador, personal_involucrado, fecha_auditoria):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = f'INFORME_AUDITORIA_{nombre_empresa}_{timestamp}.docx'
    
    doc = Document()

    # Añadir contenido al informe
    doc.add_heading(f'INFORME AUDITORIA {nombre_empresa.upper()}', 0)

    # Índice
    doc.add_paragraph('Índice')
    indice = [
        "1. Resumen Ejecutivo",
        "2. Antecedentes",
        "3. Alegaciones y Evaluación Inicial",
        "4. Alcance",
        "5. Metodología",
        "6. Transacciones Revisadas y Hallazgos",
        "7. Pruebas Realizadas",
        "8. Resumen de Pruebas Realizadas",
        "9. Cómo se Perpetró el Fraude",
        "10. Identificación de los Sospechosos",
        "11. Cuantificación de la Pérdida",
        "12. Sugerencias de Mejora de los Controles Internos",
        "13. Presencia del Auditor Forense en Procedimientos Judiciales",
        "14. Anexos",
        "15. Historial de Clientes Evaluados",
        "16. Riesgos Potenciales de Fraude Adicionales",
        "17. Patrones Inusuales",
        "18. Señales de Advertencia y Métodos de Robo",
        "19. Papeles de Trabajo"
    ]
    for item in indice:
        doc.add_paragraph(item)

    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        f"Este informe detalla los resultados de la auditoría forense realizada en '{nombre_empresa}' durante el año {fecha_auditoria.year}, en respuesta a sospechas de fraude por parte de {nombre_fraudador}. "
        "Se identificaron discrepancias significativas entre los pagos de los clientes y los registros contables, lo que sugiere la posibilidad de que algunos miembros del personal "
        f"de {personal_involucrado} estén desviando temporalmente fondos antes de registrarlos oficialmente.\n"
        "La investigación reveló debilidades en los controles internos de la empresa, que pudieron haber facilitado estas actividades fraudulentas."
    )

    # Antecedentes
    doc.add_heading('2. Antecedentes', level=1)
    doc.add_paragraph(
        f"'{nombre_empresa}' es una empresa con una amplia cartera de clientes. Recientemente, la gerencia notó discrepancias entre los pagos recibidos de los clientes y los registros contables oficiales. "
        "Estas discrepancias, junto con denuncias internas, llevaron a la sospecha de que el personal involucrado podría estar cometiendo actividades fraudulentas."
    )

    # Alegaciones y Evaluación Inicial
    doc.add_heading('3. Alegaciones y Evaluación Inicial', level=1)
    doc.add_paragraph(
        f"'{nombre_empresa}.' recibió múltiples informes que indicaban irregularidades en los cobros realizados por el personal de cobranzas. "
        "Se alegó que algunos pagos de clientes no coincidían con los registros contables y que los depósitos en las cuentas bancarias de la empresa se realizaban con retraso. "
        "Ante estas alegaciones, se decidió iniciar una auditoría forense para determinar la veracidad de las acusaciones y la magnitud del fraude."
    )

    # Alcance
    doc.add_heading('4. Alcance', level=1)
    doc.add_paragraph(
        "El alcance de la auditoría forense abarcó:\n"
        "• Revisión de los registros de cobros y depósitos bancarios de los últimos 12 meses.\n"
        "• Entrevistas con todo el personal de cobranzas.\n"
        "• Análisis detallado del flujo de efectivo para detectar irregularidades.\n"
        "• Evaluación de los procedimientos de supervisión interna y controles antifraude."
    )

    # Metodología
    doc.add_heading('5. Metodología', level=1)
    doc.add_paragraph(
        "La metodología aplicada incluyó:\n"
        "• Revisión documental: Análisis de los registros de cobros y extractos bancarios.\n"
        "• Entrevistas estructuradas: Realización de entrevistas con el personal clave, enfocadas en el manejo de cobros.\n"
        "• Análisis financiero: Implementación de técnicas de análisis de flujo de efectivo para identificar patrones inusuales.\n"
        "• Evaluación de controles: Revisión de la efectividad de los controles internos y procedimientos de supervisión."
    )

    # Transacciones Revisadas y Hallazgos
    doc.add_heading('6. Transacciones Revisadas y Hallazgos', level=1)
    doc.add_paragraph(
        "Desviación Temporal de Pagos de Clientes\n"
        "Hallazgo: Se detectaron múltiples casos donde los pagos de los clientes fueron desviados temporalmente antes de ser registrados en los sistemas contables. "
        "Estos fondos eran depositados en cuentas personales del personal de cobranzas y luego transferidos a las cuentas de la empresa, generando un retraso en los registros oficiales.\n"
        "Impacto: La desviación temporal de fondos afectó la integridad de los registros contables y pudo haber ocasionado pérdidas financieras no cuantificadas en intereses y penalidades.\n\n"
        
        "Falta de Coincidencia en Registros Contables\n"
        "Hallazgo: Hubo una falta de coincidencia recurrente entre los montos registrados como recibidos en los libros contables y los depósitos bancarios. En algunos casos, los montos registrados eran menores que los pagos reales efectuados por los clientes.\n"
        "Impacto: Esta discrepancia sugiere manipulación en los registros contables para encubrir la retención temporal de fondos. Esto pone en duda la exactitud de los estados financieros de la empresa.\n\n"
        
        "Debilidades en la Supervisión de Depósitos Bancarios\n"
        "Hallazgo: La revisión reveló que no existían controles adecuados para supervisar los depósitos bancarios realizados por el personal de cobranzas. Los depósitos se realizaban sin supervisión directa ni verificación independiente, lo que facilitó el jineteo de cobranzas.\n"
        "Impacto: La falta de supervisión permitió que el fraude ocurriera sin ser detectado durante un período prolongado, aumentando el riesgo de pérdidas financieras para la empresa.\n\n"
        
        "Manipulación de Registros de Cobranzas\n"
        "Hallazgo: Se identificaron varios casos de manipulación de registros de cobranzas, donde los pagos de clientes eran registrados en fechas posteriores a las de los depósitos bancarios. Esta manipulación permitió a los cobradores retener temporalmente los fondos antes de registrarlos oficialmente.\n"
        "Impacto: La manipulación de registros distorsionó la realidad financiera de la empresa, afectando su capacidad para tomar decisiones basadas en datos precisos y confiables.\n\n"
        
        "Evaluación de Carteras Vencidas a más de 90 Días\n"
        "Hallazgo: Las carteras de clientes con deudas vencidas a más de 90 días presentaron patrones de comportamiento sospechosos, incluyendo retrasos inexplicables en los registros de pagos y discrepancias en las fechas de depósito. Estas cuentas son especialmente vulnerables al jineteo de cobranzas, ya que la falta de seguimiento adecuado permite a los cobradores manipular los registros y retener temporalmente los fondos.\n"
        "Impacto: La falta de controles y supervisión en estas carteras puede resultar en pérdidas financieras significativas para la empresa y aumentar el riesgo de fraude. Identificar patrones de comportamiento en estas cuentas es esencial para prevenir el desvío de fondos y asegurar la integridad de los registros financieros."
    )

    # Pruebas Realizadas
    doc.add_heading('7. Pruebas Realizadas', level=1)
    doc.add_paragraph(
        "Revisión de Registros de Cobros y Depósitos Bancarios\n"
        "Objetivo: Verificar la correspondencia entre los pagos registrados y los depósitos bancarios realizados.\n"
        "Procedimiento: Se compararon los registros de cobros con los extractos bancarios para identificar discrepancias en montos y fechas.\n"
        "Resultado: Se encontraron múltiples discrepancias que confirmaron la desviación temporal de fondos.\n\n"

        "Entrevistas con el Personal de Cobranzas\n"
        "Objetivo: Obtener información sobre los procedimientos de manejo de cobros y detectar irregularidades.\n"
        "Procedimiento: Se entrevistó a todo el personal de cobranzas, con un enfoque particular en las actividades del personal sospechoso.\n"
        "Resultado: Las entrevistas revelaron que algunos cobradores tenían acceso no controlado a los fondos, y que no había supervisión adecuada de sus actividades.\n\n"

        "Análisis de Flujo de Efectivo\n"
        "Objetivo: Identificar patrones inusuales en el flujo de efectivo que pudieran indicar retención temporal de pagos.\n"
        "Procedimiento: Se analizó el flujo de efectivo durante el último año para detectar anomalías.\n"
        "Resultado: Se observaron patrones que sugerían la retención de fondos antes de su depósito en las cuentas de la empresa.\n\n"

        "Revisión de Permisos y Roles\n"
        "Objetivo: Evaluar si los permisos y roles asignados al personal de cobranzas eran adecuados y seguros.\n"
        "Procedimiento: Se revisaron los permisos de acceso al sistema contable y se compararon con las funciones reales del personal.\n"
        "Resultado: Se descubrió que algunos cobradores tenían permisos excesivos que les permitían manipular registros sin supervisión.\n\n"

        "Análisis de Documentos de Pago\n"
        "Objetivo: Verificar la autenticidad y exactitud de los documentos de pago procesados por el personal de cobranzas.\n"
        "Procedimiento: Se analizaron los recibos de pago y comprobantes bancarios para identificar inconsistencias.\n"
        "Resultado: Se encontraron documentos con fechas alteradas y montos incorrectos.\n\n"

        "Revisión de Procedimientos de Supervisión Interna\n"
        "Objetivo: Evaluar la efectividad de los procedimientos de supervisión interna en la prevención y detección de fraudes.\n"
        "Procedimiento: Se revisaron los procedimientos actuales y se compararon con las mejores prácticas de la industria.\n"
        "Resultado: Se identificaron varias áreas de mejora en los procedimientos de supervisión, que actualmente son insuficientes para prevenir fraudes.\n\n"

        "Análisis de Antigüedad de la Cartera de Clientes\n"
        "Objetivo: Evaluar la antigüedad de las cuentas por cobrar para identificar posibles áreas de riesgo de fraude, especialmente en carteras vencidas a más de 90 días.\n"
        "Procedimiento:\n"
        "• Se realizó un análisis detallado de la antigüedad de la cartera de clientes, clasificando las cuentas por cobrar según el tiempo transcurrido desde su fecha de vencimiento.\n"
        "• Se identificaron las cuentas con vencimientos superiores a 30, 60, 90 días y más, con un enfoque especial en aquellas con más de 90 días de antigüedad.\n"
        "• Se revisaron los registros de pago y se compararon con las fechas de depósito en las cuentas bancarias de la empresa.\n"
        "Resultado:\n"
        "• El análisis reveló que un número significativo de cuentas con antigüedad superior a 90 días presentaba discrepancias entre los registros de cobro y los depósitos bancarios. Estos hallazgos sugieren que estas cuentas están siendo manipuladas para retener temporalmente los fondos antes de ser registrados oficialmente."
    )

    # Resumen de Pruebas Realizadas
    doc.add_heading('8. Resumen de Pruebas Realizadas', level=1)
    doc.add_paragraph(
        f"Las pruebas realizadas confirmaron la existencia de debilidades significativas en los controles internos de {nombre_empresa}"
        "Estas debilidades permitieron a algunos miembros del personal de cobranzas desviar temporalmente los pagos de clientes, manipular registros contables y retrasar los depósitos bancarios. "
        "La falta de supervisión y controles efectivos fue un factor clave que facilitó la ocurrencia del fraude."
    )

    # Cómo se Perpetró el Fraude
    doc.add_heading('9. Cómo se Perpetró el Fraude', level=1)
    doc.add_paragraph(
        "El fraude de jineteo de cobranzas posiblemente está siendo perpetrado por algunos miembros del personal de cobranzas, quienes aprovechan las debilidades en los controles internos para desviar temporalmente los pagos de clientes. "
        "Estas actividades fraudulentas son particularmente evidentes en las carteras vencidas a más de 90 días, donde los cobradores manipulan los registros y retrasan los depósitos en las cuentas de la empresa, lo que permite que los fondos sean retenidos temporalmente sin detección inmediata."
    )

    # Identificación de los Sospechosos
    doc.add_heading('10. Identificación de los Sospechosos', level=1)
    doc.add_paragraph(
        f"El principal sospechoso identificado es {nombre_fraudador}, cobrador de '{nombre_empresa}'. Las pruebas indican que {nombre_fraudador} tenía acceso no controlado a los fondos y la capacidad de manipular los registros contables. "
        "No se encontraron evidencias de la participación de otros empleados en este fraude."
    )

    # Cuantificación de la Pérdida
    doc.add_heading('11. Cuantificación de la Pérdida', level=1)
    perdida_total = pagos_vencidos_90_dias['SALDO'].sum()
    doc.add_paragraph(
        f"Estimación de la Pérdida: La pérdida financiera exacta aún no se ha determinado, pero se estima que podría alcanzar los ${perdida_total:,.2f}, considerando el valor de los pagos desviados temporalmente, los intereses perdidos y las posibles sanciones por incumplimiento de obligaciones fiscales."
    )

    # Sugerencias de Mejora de los Controles Internos
    doc.add_heading('12. Sugerencias de Mejora de los Controles Internos', level=1)
    doc.add_paragraph(
        "1. Mejora en la Segregación de Funciones: Implementar una segregación estricta de funciones para evitar que una sola persona tenga control total sobre los cobros y depósitos.\n"
        "2. Supervisión de Depósitos: Establecer un proceso de verificación independiente para todos los depósitos bancarios realizados por el personal de cobranzas.\n"
        "3. Automatización de Registros: Implementar un sistema automatizado para el registro de cobros que incluya alertas automáticas para discrepancias en montos y fechas.\n"
        "4. Auditorías Periódicas: Realizar auditorías internas periódicas centradas en el área de cobranzas y manejo de efectivo.\n"
        "5. Capacitación del Personal: Capacitar al personal en prácticas de gestión de riesgos y la importancia de la integridad en el manejo de fondos."
    )

    # Presencia del Auditor Forense en Procedimientos Judiciales
    doc.add_heading('13. Presencia del Auditor Forense en Procedimientos Judiciales', level=1)
    doc.add_paragraph(
        "El auditor forense deberá estar presente durante los procedimientos judiciales para presentar y explicar las pruebas recopiladas. "
        "Esto incluye demostrar cómo se identificó el fraude, la metodología utilizada, y la autenticidad de las pruebas. "
        "La presencia del auditor es crucial para respaldar la acusación contra los responsables y asegurar que la justicia se imparta adecuadamente."
    )

    # Anexos
    doc.add_heading('14. Anexos', level=1)
    doc.add_paragraph(
        "Anexo 1: Detalle de los registros de cobros y depósitos revisados.\n"
        "Anexo 2: Transcripciones de entrevistas con el personal de cobranzas.\n"
        "Anexo 3: Resultados del análisis de flujo de efectivo.\n"
        "Anexo 4: Documentación sobre la revisión de permisos y roles.\n"
        "Anexo 5: Análisis de antigüedad de la cartera de clientes."
    )

    # Historial de Clientes Evaluados
    doc.add_heading('15. Historial de Clientes Evaluados', level=1)
    for cliente in historial_clientes:
        doc.add_paragraph(cliente, style='Normal')

    # Riesgos Potenciales de Fraude Adicionales
    doc.add_heading('16. Riesgos Potenciales de Fraude Adicionales', level=1)
    doc.add_paragraph(
        "Durante la auditoría, se identificaron otros posibles riesgos de fraude que requieren atención:"
    )
    riesgos = [
        "Retención de pagos por períodos prolongados antes de su registro oficial.",
        "Ajustes frecuentes en las cuentas sin justificación adecuada.",
        "Discrepancias entre las cantidades registradas y los depósitos reales.",
        "Falta de supervisión efectiva en los procesos de autorización de pagos."
    ]
    for riesgo in riesgos:
        doc.add_paragraph(f"• {riesgo}")

    # Patrones inusuales que pueden indicar jineteo de fondos
    doc.add_heading('17. Patrones Inusuales', level=1)
    doc.add_paragraph(
        "Existen ciertos patrones inusuales que pueden indicar la presencia de jineteo de fondos en la empresa. "
        "Estos incluyen:"
    )
    patrones = [
        "Retrasos inexplicables en la contabilización de los pagos recibidos.",
        "Frecuentes ajustes a la cuenta de clientes sin una razón clara.",
        "Discrepancias entre los registros de pagos y los extractos bancarios.",
        "Pagos recibidos que no coinciden con los registros de la cuenta por cobrar.",
        "Depósitos en la cuenta bancaria que no se reflejan inmediatamente en los registros contables."
    ]
    for patron in patrones:
        doc.add_paragraph(f"• {patron}")

    # Señales de advertencia y formas en que los cobradores pueden robarse el dinero
    doc.add_heading('18. Señales de Advertencia y Métodos de Robo', level=1)
    doc.add_paragraph(
        "Los siguientes son algunas señales de advertencia y métodos comunes que los cobradores pueden utilizar para "
        "robarse el dinero:"
    )
    señales = [
        "Pagos en efectivo no depositados de inmediato en la cuenta bancaria.",
        "Frecuentes quejas de clientes sobre pagos que no han sido registrados.",
        "Falsificación de documentos para justificar faltantes en los depósitos.",
        "Retrasos en la entrega de reportes financieros.",
        "Cobradores que insisten en recibir pagos en efectivo en lugar de cheques o transferencias.",
        "Cuentas que se mantienen abiertas durante un tiempo prolongado sin ser saldadas a pesar de los pagos realizados."
    ]
    for señal in señales:
        doc.add_paragraph(f"• {señal}")

    # Papeles de Trabajo
    doc.add_heading('19. Papeles de Trabajo', level=1)
    doc.add_paragraph(
        "Para descubrir este tipo de fraude, los auditores deben realizar una serie de procedimientos detallados, incluyendo:"
    )
    procedimientos = [
        "Comparación de los registros de pagos con los extractos bancarios para asegurar que los pagos fueron depositados en tiempo y forma.",
        "Verificación de los procedimientos de autorización y registro de pagos.",
        "Análisis de patrones inusuales en los registros de pagos y depósitos.",
        "Evaluación de la segregación de funciones en el proceso de manejo de pagos."
    ]
    for procedimiento in procedimientos:
        doc.add_paragraph(f"• {procedimiento}")

    doc.add_paragraph(
        "Además, es importante realizar entrevistas y confirmar directamente con los clientes los pagos realizados y sus fechas. "
        "Esto puede ayudar a identificar discrepancias y posibles fraudes."
    )

    # Agregar imágenes de papeles de trabajo
    doc.add_heading('Imágenes de Papeles de Trabajo', level=2)
    for imagen in imagenes_papeles_paths:
        doc.add_picture(imagen, width=Inches(5.0))  # Ajusta el tamaño de la imagen según sea necesario

    # Guardar el documento de Word
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.success(f"Informe de auditoría guardado en '{file_name}'.")

    # Descargar el archivo de Word
    st.download_button(label="Descargar archivo Word", data=buffer, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Función para analizar las anomalías de la cartera
def analizar_anomalias_cartera(file):
    xls = pd.ExcelFile(file)
    pagos_vencidos_90_dias = []

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        df.columns = df.columns.str.strip()

        # Asegurarse de que las columnas necesarias estén presentes
        if 'DÍAS DE MORA' in df.columns and 'SALDO' in df.columns:
            df['SALDO'] = df['SALDO'].replace('[\$,]', '', regex=True).astype(float)
            vencidos_90_dias = df[(df['DÍAS DE MORA'] >= 90) & (df['SALDO'] > 0)]
            if not vencidos_90_dias.empty:
                pagos_vencidos_90_dias.append(vencidos_90_dias)
        else:
            st.error(f"Hoja {sheet_name} no contiene las columnas necesarias 'DÍAS DE MORA' y/o 'SALDO'.")

    if pagos_vencidos_90_dias:
        pagos_vencidos_90_dias_df = pd.concat(pagos_vencidos_90_dias, ignore_index=True)
        return pagos_vencidos_90_dias_df
    else:
        st.warning("No se encontraron pagos vencidos a más de 90 días.")
        return None

# Streamlit UI
st.title("Auditoría Forense")
st.markdown("""
**Bienvenido a la herramienta de auditoría forense**. Esta aplicación está diseñada para ayudar en la identificación y análisis de anomalías en las cuentas por cobrar de la empresa, especialmente enfocándose en la detección de fraudes como el jineteo de cobranzas.
""")

st.header("Descripción de la herramienta")
st.markdown("""
Esta herramienta permite cargar y analizar datos de pagos vencidos de clientes, generando informes detallados en Excel y Word. 
Las principales funcionalidades incluyen:

- **Análisis de carteras vencidas**: Identificación de cuentas con pagos vencidos a más de 90 días.
- **Generación de informes**: Creación de informes en formatos Excel y Word con detalles específicos sobre las transacciones revisadas y hallazgos.
- **Detección de fraudes**: Identificación de patrones inusuales que podrían indicar actividades fraudulentas, como el jineteo de cobranzas.

### Instrucciones de uso:
1. **Subir archivo Excel**: Carga el archivo de Excel con las carteras vencidas para iniciar el análisis.
2. **Subir archivo Word** (opcional): Carga un archivo de Word con el historial de clientes para incluir en el informe.
3. **Subir imágenes de papeles de trabajo**: Suba aquí los papeles de trabajo en formato JPG que desee incluir en el informe.
4. **Descargar informes**: Una vez procesados los datos, descarga los informes generados en los formatos proporcionados.
""")

# Subir archivo Excel para análisis de carteras vencidas
st.header("Subir archivo Excel")
st.markdown("Por favor, sube el archivo Excel que contiene la información de las carteras vencidas.")

# Añadir el botón de descarga del archivo de ejemplo aquí
st.markdown("Si no tienes un archivo de ejemplo, puedes descargar una plantilla de ejemplo aquí:")

try:
    with open("Plantilla Evaluacion de cartera.xlsx", "rb") as f:
        st.download_button(label="Descargar plantilla de ejemplo", data=f, file_name="Plantilla_Evaluacion_de_cartera.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
except FileNotFoundError:
    st.error("No se pudo encontrar la plantilla de ejemplo. Asegúrate de que el archivo está en la ubicación correcta.")

# Solicitar archivo Excel
file_excel = st.file_uploader("Seleccione el archivo Excel con las carteras vencidas", type=["xlsx", "xls"])

if file_excel:
    pagos_vencidos_90_dias_df = analizar_anomalias_cartera(file_excel)

    if pagos_vencidos_90_dias_df is not None:
        generar_informe_excel(pagos_vencidos_90_dias_df)

        # Subir archivo Word para historial de clientes
        st.header("Subir archivo Word")
        st.markdown("Opcional: Sube un archivo Word que contenga el historial de clientes que desees incluir en el informe final.")
        file_word = st.file_uploader("Seleccione el archivo Word con el historial de clientes", type=["docx"])

        if file_word:
            historial_clientes = extraer_historial_clientes(file_word)
        else:
            historial_clientes = []

        # Subir imágenes de papeles de trabajo
        st.header("Subir imágenes de Papeles de Trabajo")
        st.markdown("Sube las imágenes en formato JPG que deseas incluir en los papeles de trabajo del informe final.")
        imagenes_papeles = st.file_uploader("Seleccione las imágenes de papeles de trabajo", type=["jpg", "jpeg"], accept_multiple_files=True)

        if imagenes_papeles:
            # Guardar las imágenes temporalmente
            imagenes_papeles_paths = []
            for img in imagenes_papeles:
                image_path = f"/tmp/{img.name}"
                with open(image_path, "wb") as f:
                    f.write(img.getbuffer())
                imagenes_papeles_paths.append(image_path)
        else:
            imagenes_papeles_paths = []

        # Formulario para ingresar datos adicionales
        st.header("Formulario de datos de la auditoría")
        nombre_empresa = st.text_input("Nombre de la empresa")
        nombre_fraudador = st.text_input("Nombre del posible defraudador")
        jefe_personal_involucrado = st.text_input("Jefe del Personal involucrado en el manejo de fondos")
        fecha_auditoria = st.date_input("Fecha de la auditoría")

        if st.button("Generar Informe de Auditoría"):
            if nombre_empresa and nombre_fraudador and jefe_personal_involucrado and fecha_auditoria:
                generar_informe_word(pagos_vencidos_90_dias_df, historial_clientes, imagenes_papeles_paths, nombre_empresa, nombre_fraudador, jefe_personal_involucrado, fecha_auditoria)
            else:
                st.error("Por favor, complete todos los campos del formulario antes de generar el informe.")
